from __future__ import annotations

from io import BytesIO

from docx import Document


def build_report_docx(title: str, sections: list[tuple[str, list[str]]]) -> bytes:
    """Jamoa hisobotini .docx sifatida quradi va bytes qaytaradi.

    sections: [(bo'lim sarlavhasi, [qatorlar]), ...]"""
    doc = Document()
    doc.add_heading(title, level=0)
    for header, lines in sections:
        doc.add_heading(header, level=1)
        if not lines:
            doc.add_paragraph("—")
            continue
        for line in lines:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
